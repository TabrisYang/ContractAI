"""
合約去識別化引擎
七層偵測：Regex → 合約上下文 → 字典比對 → 商業機密 → NER → TF-IDF → LLM（最後防線）
"""
import asyncio
import json
import re
import time
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

import joblib
import numpy as np
import spacy
from docx import Document
from loguru import logger

from ..models.schemas import MaskingMethod
from .config import settings

logger.add(
    settings.LOG_DIR / "deidentifier.log",
    rotation="500 MB",
    retention="30 days",
    level="INFO",
    format="{time:YYYY-MM-DD HH:mm:ss} | {level} | {message}",
)

MASK_MAP = {
    "ID": "[身分證字號]",
    "PHONE": "[電話]",
    "TAX_ID": "[統一編號]",
    "EMAIL": "[電子郵件]",
    "DATE": "[日期]",
    "MONEY": "[金額]",
    "PERSON": "[姓名]",
    "ORG": "[機構名稱]",
    "GPE": "[地名]",
    "LOC": "[地點]",
    "ADDRESS": "[地址]",
    "PERCENT": "[百分比]",
    "RARE_TERM": "[罕見詞]",
    "LLM_PII": "[敏感資訊]",
    "BANK_ACCOUNT": "[銀行帳號]",
    "BANK": "[銀行]",
    "FAX": "[傳真]",
    "BRAND": "[品牌名稱]",
    "PROJECT_NAME": "[專案名稱]",
}

# ── 合約上下文關鍵字 → 後面接 PII ────────────────────────────────
CONTEXT_PATTERNS: Dict[str, List[str]] = {
    "PERSON": [
        (
            r'(?:甲[\s　]*方[\s　]*代[\s　]*表[\s　]*人|乙[\s　]*方[\s　]*代[\s　]*表[\s　]*人'
            r'|丙[\s　]*方[\s　]*代[\s　]*表[\s　]*人'
            r'|代[\s　]*表[\s　]*人|法[\s　]*定[\s　]*代[\s　]*理[\s　]*人|負[\s　]*責[\s　]*人'
            r'|連帶保證人|共同發票人|立[\s　]*契[\s　]*約[\s　]*書[\s　]*人|立書人|立約人|立合約書人'
            r'|聯絡人|聯絡窗口|專案經理|經辦人|承辦人|見證人|保證人'
            r'|被保險人|要保人|受益人|簽署人|授權代表|經理人'
            r'|甲[\s　]*方[\s　]*簽[\s　]*章|乙[\s　]*方[\s　]*簽[\s　]*章'
            r'|丙[\s　]*方[\s　]*簽[\s　]*章)[\s　：:]+\s*'
            r'([^\s　，。、；\n\(\)（）\[\]]{2,5})'
        ),
    ],
    "ORG": [
        (
            r'(?:甲[\s　]*方|乙[\s　]*方|丙[\s　]*方|委託人|受託人|出借人|借款人|承攬人|定作人'
            r'|買方|賣方|出租人|承租人|授權人|被授權人|債權人|債務人'
            r'|委任人|受任人|供應商|廠商|合作方)[\s　：:]+\s*'
            r'([^\s　，。\n\[\]]{2,25}?'
            r'(?:股份有限公司|有限公司|公司|企業社|工作室|事務所|基金會|協會|商行))'
        ),
    ],
    "PHONE": [
        (
            r'(?:電話|Tel|TEL|tel|聯絡電話|手機|行動電話|連絡電話|公司電話|辦公電話)[\s：:]+\s*'
            r'([\d\-\(\)\s\+]{7,20})'
        ),
    ],
    "FAX": [
        (
            r'(?:傳真|Fax|FAX|fax)[\s：:]+\s*'
            r'([\d\-\(\)\s\+]{7,20})'
        ),
    ],
    "ADDRESS": [
        (
            r'(?:地[\s　]*址|住[\s　]*址|戶籍地址|通訊地址|營業地址|公司地址|住所|營業所|所在地)[\s　：:]+\s*'
            r'([^\n]{5,60}?)(?=[。\n]|$)'
        ),
    ],
    "BANK_ACCOUNT": [
        r'(?:帳號|帳戶|銀行帳號|匯款帳號|收款帳號|帳戶號碼|帳戶帳號)[\s：:]+\s*([\d\-\s]{5,25})',
        r'(?:戶名)[\s：:]+\s*([^\n，。；]{2,25})',
    ],
    "BANK": [
        r'(?:銀行|匯款銀行|收款銀行)[\s：:]+\s*([^\n，。；]{2,20})',
    ],
    "EMAIL": [
        r'(?:Email|e-?mail|E-?MAIL|電子郵件|信箱|電郵)[\s：:]+\s*([^\s，。\n]+@[^\s，。\n]+)',
    ],
}

# ── 字典：台灣銀行名稱 ────────────────────────────────────────
BANK_NAMES = [
    "臺灣銀行", "台灣銀行", "土地銀行", "合作金庫", "第一銀行", "華南銀行",
    "彰化銀行", "兆豐銀行", "台灣企銀", "國泰世華", "台北富邦",
    "高雄銀行", "台中銀行", "中國信託", "上海商銀", "遠東銀行",
    "元大銀行", "永豐銀行", "玉山銀行", "凱基銀行", "星展銀行",
    "台新銀行", "日盛銀行", "安泰銀行", "聯邦銀行", "陽信銀行",
    "華泰銀行", "瑞興銀行", "王道銀行", "樂天銀行",
]

# ── 字典：百家姓（台灣前 100 大姓） ────────────────────────────
SURNAMES = (
    "陳林黃張李王吳劉蔡楊許鄭謝洪郭邱曾廖賴徐周葉蘇莊呂江何蕭羅高潘簡朱鍾游彭詹"
    "胡施沈余趙盧梁顏柯翁魏孫戴範方宋鄧杜傅侯曹薛丁卓馬阮董溫蒲藍石連褚紀蔣"
    "童歐尤田巫涂鄒康池白湯姚韓龔譚賈俞段雷錢"
)

# 合約中不該被當成品牌遮罩的法律常用英文詞
LEGAL_ENGLISH_WHITELIST = {
    "Article", "Section", "Party", "Parties", "Agreement", "Contract",
    "Appendix", "Schedule", "Exhibit", "Clause", "Term", "Terms",
    "The", "This", "That", "For", "And", "With", "From", "Date",
    "Service", "Services", "Project", "Work", "Scope", "Payment",
    "Price", "Total", "Amount", "Period", "Notice", "Confidential",
    "NTD", "USD", "TWD", "SOW", "NDA", "PO", "RFP", "RFQ",
    "PDF", "DOCX", "URL", "CEO", "CFO", "CTO", "COO", "PM",
    "KOL", "KPI", "ROI", "SOP", "FAQ", "API", "CRM", "ERP",
    "Inc", "Ltd", "Corp", "LLC", "Co",
}


class DocumentDeidentifier:
    """合約去識別化主要類別"""

    def __init__(self):
        self.nlp = None
        self.tfidf_vectorizer = None
        self._load_models()

        # ── 第 1 層：基礎 Regex ──
        self.patterns = {
            "ID": r'[A-Z][12]\d{8}',
            "PHONE": (
                r'(?:(?:\+886|0)[-\s]?)?(?:9\d{2}[-\s]?\d{3}[-\s]?\d{3})'  # 手機
                r'|0[2-9][-\s]?\d{3,4}[-\s]?\d{4}'                          # 市話 0X-XXXX-XXXX
                r'|\(\d{2,3}\)\s?\d{3,4}[-\s]?\d{4}'                        # (0X)XXXX-XXXX
            ),
            "TAX_ID": r'(?:統[\s　]*一[\s　]*編[\s　]*號|統[\s　]*編)[\s　：:]+\s*(\d{8})',
            "EMAIL": r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}',
            "DATE": r'\d{2,4}[-/年]\d{1,2}[-/月]\d{1,2}[日]?',
            "MONEY": (
                r'(?:NT\$|NTD|新臺幣|新台幣|TWD|美金|USD|\$)\s*\d{1,3}(?:,\d{3})*(?:\.\d+)?'
                r'|\d{1,3}(?:,\d{3})*(?:\.\d+)?\s*(?:元整|元|萬元|億元|千元)'
            ),
            "ADDRESS": (
                r'(?:(?:台|臺)[北中南東灣]|新北|桃園|新竹|苗栗|彰化|南投|雲林|嘉義|屏東|宜蘭|花蓮|台東|澎湖|金門|連江)'
                r'(?:市|縣)'
                r'(?:[\u4e00-\u9fff]{1,4}(?:區|鄉|鎮|市))?'
                r'[^\n，。；]{0,30}?(?:路|街|大道|巷|弄|號|樓)[^\n，。；]{0,15}'
            ),
        }

    def _load_models(self):
        logger.info("載入模型...")
        try:
            self.nlp = spacy.load(settings.SPACY_MODEL)
            logger.info(f"已載入 spaCy 模型：{settings.SPACY_MODEL}")
        except OSError:
            logger.warning(f"找不到 {settings.SPACY_MODEL}，正在下載...")
            spacy.cli.download(settings.SPACY_MODEL)
            self.nlp = spacy.load(settings.SPACY_MODEL)

        if settings.TFIDF_MODEL_PATH.exists():
            self.tfidf_vectorizer = joblib.load(settings.TFIDF_MODEL_PATH)
            logger.info("已載入 TF-IDF 模型")
        else:
            logger.warning("找不到 TF-IDF 模型，跳過罕見詞偵測")

    def process(
        self,
        input_path: str,
        output_dir: Path,
        job_id: str,
        callback: Optional[Callable[[float, str], None]] = None,
        **options,
    ) -> Dict[str, Any]:
        start_time = time.time()

        def progress(p: float, msg: str):
            if callback:
                callback(p, msg)
            logger.info(f"[{job_id}] {p:.0f}% - {msg}")

        progress(0, "讀取文件")
        doc = Document(input_path)
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)

        raw_methods = options.get(
            "mask_methods",
            [MaskingMethod.REGEX, MaskingMethod.NER, MaskingMethod.TFIDF],
        )
        mask_methods = []
        for m in raw_methods:
            try:
                mask_methods.append(MaskingMethod(m) if isinstance(m, str) else m)
            except ValueError:
                pass

        all_entities: List[Dict] = []
        stats: Dict[str, int] = {}

        # ── 第 1 層：Regex 基礎模式 ──
        if MaskingMethod.REGEX in mask_methods:
            progress(10, "Regex 模式偵測")
            ents = self._apply_regex_masking(full_text)
            all_entities.extend(ents)
            stats["regex"] = len(ents)

        # ── 第 2 層：合約上下文偵測（新增） ──
        progress(20, "合約上下文偵測")
        ents = self._apply_context_masking(full_text)
        all_entities.extend(ents)
        stats["context"] = len(ents)

        # ── 第 3 層：字典比對（新增） ──
        progress(35, "字典比對偵測")
        ents = self._apply_dictionary_masking(full_text)
        all_entities.extend(ents)
        stats["dictionary"] = len(ents)

        # ── 第 4 層：商業機密偵測（新增） ──
        progress(45, "商業機密偵測")
        ents = self._apply_business_masking(full_text)
        all_entities.extend(ents)
        stats["business"] = len(ents)

        # ── 第 5 層：NER ──
        if MaskingMethod.NER in mask_methods and self.nlp:
            progress(55, "NER 命名實體辨識")
            ents = self._apply_ner_masking(full_text)
            all_entities.extend(ents)
            stats["ner"] = len(ents)

        # ── 第 6 層：TF-IDF ──
        if MaskingMethod.TFIDF in mask_methods and self.tfidf_vectorizer:
            progress(65, "TF-IDF 罕見詞偵測")
            ents = self._apply_tfidf_masking(full_text)
            all_entities.extend(ents)
            stats["tfidf"] = len(ents)

        # ── 傳播：已知實體的其他出現位置 ──
        progress(70, "實體傳播比對")
        ents = self._propagate_known_entities(full_text, all_entities)
        all_entities.extend(ents)
        stats["propagation"] = len(ents)

        # ── 第 7 層：LLM 最後防線 ──
        llm_config = options.get("llm_config")
        if MaskingMethod.LLM in mask_methods and llm_config:
            progress(75, "LLM 輔助偵測")
            try:
                loop = asyncio.new_event_loop()
                ents = loop.run_until_complete(
                    self._apply_llm_masking(full_text, all_entities, llm_config)
                )
                loop.close()
                all_entities.extend(ents)
                stats["llm"] = len(ents)
            except Exception as e:
                logger.warning(f"LLM 輔助偵測失敗（跳過）：{e}")

        progress(85, "合併重疊實體")
        merged = self._merge_overlapping_entities(all_entities)

        entities_by_type: Dict[str, int] = {}
        for e in merged:
            t = e["entity_type"]
            entities_by_type[t] = entities_by_type.get(t, 0) + 1

        progress(90, "套用遮罩並儲存文件")
        output_path = self._apply_masks_to_document(
            doc=doc,
            entities=merged,
            full_text=full_text,
            output_dir=output_dir,
            job_id=job_id,
        )

        elapsed = round(time.time() - start_time, 2)
        progress(100, "處理完成")

        return {
            "status": "completed",
            "job_id": job_id,
            "input_path": input_path,
            "output_path": str(output_path),
            "analysis": {
                "total_characters": len(full_text),
                "total_entities": len(merged),
                "entities_by_type": entities_by_type,
                "processing_time": elapsed,
                "masking_stats": stats,
            },
        }

    # ── 第 1 層：Regex ──────────────────────────────────────────

    def _apply_regex_masking(self, text: str) -> List[Dict]:
        entities = []
        for entity_type, pattern in self.patterns.items():
            for m in re.finditer(pattern, text):
                # TAX_ID 使用 capture group
                matched_text = m.group(1) if m.lastindex else m.group()
                start = m.start(1) if m.lastindex else m.start()
                end = m.end(1) if m.lastindex else m.end()
                entities.append({
                    "text": matched_text,
                    "entity_type": entity_type,
                    "start_pos": start,
                    "end_pos": end,
                    "method": "regex",
                    "confidence": 1.0,
                })
        return entities

    # ── 第 2 層：合約上下文偵測 ──────────────────────────────────

    def _apply_context_masking(self, text: str) -> List[Dict]:
        """利用合約固定格式的關鍵字，偵測後面接的 PII"""
        entities = []
        for entity_type, patterns in CONTEXT_PATTERNS.items():
            for pattern in patterns:
                for m in re.finditer(pattern, text):
                    if m.lastindex:
                        matched = m.group(1).strip()
                        start = m.start(1)
                        end = m.start(1) + len(matched)
                    else:
                        matched = m.group().strip()
                        start = m.start()
                        end = m.end()
                    if len(matched) < 2:
                        continue
                    entities.append({
                        "text": matched,
                        "entity_type": entity_type,
                        "start_pos": start,
                        "end_pos": end,
                        "method": "context",
                        "confidence": 0.95,
                    })
        return entities

    # ── 第 3 層：字典比對 ────────────────────────────────────────

    def _apply_dictionary_masking(self, text: str) -> List[Dict]:
        """字典比對：公司名稱、銀行名稱、中文人名"""
        entities = []

        # a. 公司名稱（中文字 + 公司後綴）
        org_pattern = r'[\u4e00-\u9fff]{2,15}(?:股份有限公司|有限公司|企業社|工作室|事務所|基金會|協會|商行)'
        for m in re.finditer(org_pattern, text):
            entities.append({
                "text": m.group(),
                "entity_type": "ORG",
                "start_pos": m.start(),
                "end_pos": m.end(),
                "method": "dictionary",
                "confidence": 0.95,
            })

        # b. 銀行名稱（+ 可選分行名）
        bank_pattern = (
            r'(?:' + '|'.join(re.escape(b) for b in BANK_NAMES) + r')'
            r'(?:商業銀行)?'
            r'(?:[\u4e00-\u9fff]{1,5}分行)?'
        )
        for m in re.finditer(bank_pattern, text):
            entities.append({
                "text": m.group(),
                "entity_type": "BANK",
                "start_pos": m.start(),
                "end_pos": m.end(),
                "method": "dictionary",
                "confidence": 0.95,
            })

        # c. 中文人名（百家姓 + 1-3 字名）
        # 只在合理的上下文中匹配：前後是標點、空白、冒號、行首行尾
        surname_chars = '[' + SURNAMES + ']'
        name_pattern = (
            r'(?:^|(?<=[\s：:，。、；）\)\n]))'
            + surname_chars
            + r'[\u4e00-\u9fff]{1,3}'
            + r'(?=$|[\s，。、；（\(\n：:])'
        )
        for m in re.finditer(name_pattern, text, re.MULTILINE):
            name = m.group().strip()
            # 排除常見非人名的中文詞（2字且為常用詞的跳過）
            if len(name) == 2 and name in _COMMON_TWO_CHAR_WORDS:
                continue
            if len(name) < 2 or len(name) > 4:
                continue
            entities.append({
                "text": name,
                "entity_type": "PERSON",
                "start_pos": m.start(),
                "end_pos": m.start() + len(name),
                "method": "dictionary",
                "confidence": 0.8,
            })

        return entities

    # ── 第 4 層：商業機密偵測 ─────────────────────────────────────

    def _apply_business_masking(self, text: str) -> List[Dict]:
        """偵測品牌名稱、專案名稱等商業機密"""
        entities = []

        # a. 英文品牌名（2+ 個英文字母開頭的詞，排除法律常用詞）
        for m in re.finditer(r'\b[A-Z][A-Za-z]{2,}(?:\s+[A-Z][A-Za-z]*)*\b', text):
            brand = m.group().strip()
            # 排除法律/通用英文詞
            words = brand.split()
            if all(w in LEGAL_ENGLISH_WHITELIST for w in words):
                continue
            if len(brand) <= 3:  # 太短的跳過（可能是縮寫）
                continue
            entities.append({
                "text": brand,
                "entity_type": "BRAND",
                "start_pos": m.start(),
                "end_pos": m.end(),
                "method": "business",
                "confidence": 0.85,
            })

        # b. 專案名稱（「」引號內 3-50 字的內容）
        for m in re.finditer(r'[「][^」]{3,50}[」]', text):
            content = m.group()
            # 排除純法律條文引用（如「本合約」「甲方」等短詞）
            inner = content[1:-1]
            if len(inner) <= 4 and not re.search(r'[A-Za-z\d]', inner):
                continue
            entities.append({
                "text": content,
                "entity_type": "PROJECT_NAME",
                "start_pos": m.start(),
                "end_pos": m.end(),
                "method": "business",
                "confidence": 0.85,
            })

        # c. 銀行代碼 + 帳號（3碼-5~14碼）
        for m in re.finditer(r'\b(\d{3})[-\s](\d{5,14})\b', text):
            entities.append({
                "text": m.group(),
                "entity_type": "BANK_ACCOUNT",
                "start_pos": m.start(),
                "end_pos": m.end(),
                "method": "business",
                "confidence": 0.9,
            })

        return entities

    # ── 實體傳播：同一實體的其他出現位置 ─────────────────────────

    @staticmethod
    def _propagate_known_entities(text: str, existing: List[Dict]) -> List[Dict]:
        """找出已偵測到的 PERSON/ORG 實體在文本中的所有其他出現位置"""
        known_positions = set()
        for e in existing:
            known_positions.update(range(e["start_pos"], e["end_pos"]))

        # 收集已知的人名和機構名
        names_to_propagate: Dict[str, str] = {}
        for e in existing:
            if e["entity_type"] in ("PERSON", "ORG", "BANK") and len(e["text"]) >= 2:
                names_to_propagate[e["text"]] = e["entity_type"]

        entities = []
        for name, etype in names_to_propagate.items():
            for m in re.finditer(re.escape(name), text):
                # 跳過已經被偵測到的位置
                if m.start() in known_positions:
                    continue
                entities.append({
                    "text": name,
                    "entity_type": etype,
                    "start_pos": m.start(),
                    "end_pos": m.end(),
                    "method": "propagation",
                    "confidence": 0.95,
                })
        return entities

    # ── 第 5 層：NER ─────────────────────────────────────────────

    def _apply_ner_masking(self, text: str) -> List[Dict]:
        if not self.nlp:
            return []
        entities = []
        chunk_size = 80000
        for i in range(0, len(text), chunk_size):
            chunk = text[i: i + chunk_size]
            doc = self.nlp(chunk)
            for ent in doc.ents:
                if ent.label_ in ("PERSON", "ORG", "GPE", "LOC", "DATE", "MONEY", "PERCENT"):
                    entities.append({
                        "text": ent.text,
                        "entity_type": ent.label_,
                        "start_pos": i + ent.start_char,
                        "end_pos": i + ent.end_char,
                        "method": "ner",
                        "confidence": 0.9,
                    })
        return entities

    # ── 第 6 層：TF-IDF ──────────────────────────────────────────

    def _apply_tfidf_masking(self, text: str) -> List[Dict]:
        if not self.tfidf_vectorizer:
            return []
        entities = []
        try:
            words = text.split()
            tfidf_scores = self.tfidf_vectorizer.transform([" ".join(words)])
            feature_names = self.tfidf_vectorizer.get_feature_names_out()
            n_top = max(1, int(len(feature_names) * settings.RARE_TERM_THRESHOLD))
            top_indices = np.argsort(tfidf_scores.toarray()[0])[:n_top]
            rare_terms = {feature_names[i] for i in top_indices}
            for term in rare_terms:
                for m in re.finditer(r'\b' + re.escape(term) + r'\b', text):
                    entities.append({
                        "text": m.group(),
                        "entity_type": "RARE_TERM",
                        "start_pos": m.start(),
                        "end_pos": m.end(),
                        "method": "tfidf",
                        "confidence": 0.7,
                    })
        except Exception as e:
            logger.error(f"TF-IDF 錯誤：{e}")
        return entities

    # ── 第 7 層：LLM 最後防線 ────────────────────────────────────

    async def _apply_llm_masking(
        self,
        text: str,
        existing_entities: List[Dict],
        llm_config: Dict,
    ) -> List[Dict]:
        """LLM 輔助偵測：先將已知 PII 遮罩，再送部分遮罩文字給 LLM 找漏網之魚"""
        from .llm.factory import create_provider
        from .llm.base import LLMMessage

        # 先用已偵測到的實體遮罩文字，避免將已知 PII 送到外部
        pre_masked = self._mask_text(text, existing_entities)

        # 取前 2000 字作為樣本
        sample = pre_masked[:2000]
        if not sample.strip():
            return []

        prompt = (
            "以下是一份合約的部分內容，其中已知的個人資訊已用 [類型] 遮罩。\n"
            "請找出其中**尚未被遮罩**的個人資訊（PII），例如遺漏的姓名、電話、地址、公司名稱、銀行帳號等。\n"
            '只回傳 JSON 陣列，格式：[{"text": "原文", "type": "PERSON/ORG/PHONE/ID/DATE/MONEY/ADDRESS/BANK_ACCOUNT"}]\n'
            "若全部都已遮罩完畢，回傳空陣列 []。\n\n"
            f"合約內容：\n{sample}"
        )

        provider_type = llm_config.get("provider", "openai")
        if hasattr(provider_type, "value"):
            provider_type = provider_type.value

        kwargs = {"model": llm_config.get("model", "")}
        if llm_config.get("api_key"):
            kwargs["api_key"] = llm_config["api_key"]
        if llm_config.get("base_url"):
            kwargs["base_url"] = llm_config["base_url"]

        llm = create_provider(provider_type, **kwargs)
        raw = await llm.simple_prompt(prompt)

        entities = []
        json_match = re.search(r'\[[\s\S]*?\]', raw)
        if json_match:
            try:
                for item in json.loads(json_match.group()):
                    pii_text = str(item.get("text", "")).strip()
                    pii_type = item.get("type", "LLM_PII")
                    if not pii_text:
                        continue
                    for m in re.finditer(re.escape(pii_text), text):
                        entities.append({
                            "text": pii_text,
                            "entity_type": pii_type,
                            "start_pos": m.start(),
                            "end_pos": m.end(),
                            "method": "llm",
                            "confidence": 0.85,
                        })
            except json.JSONDecodeError:
                logger.warning("LLM 輔助偵測回傳 JSON 解析失敗")
        return entities

    # ── 合併與套用遮罩 ──────────────────────────────────────────

    @staticmethod
    def _merge_overlapping_entities(entities: List[Dict]) -> List[Dict]:
        if not entities:
            return []
        sorted_ents = sorted(entities, key=lambda x: (x["start_pos"], -x["end_pos"]))
        merged = [sorted_ents[0].copy()]
        for cur in sorted_ents[1:]:
            last = merged[-1]
            if cur["start_pos"] < last["end_pos"]:
                if cur.get("confidence", 0) > last.get("confidence", 0):
                    last["entity_type"] = cur["entity_type"]
                last["end_pos"] = max(last["end_pos"], cur["end_pos"])
            else:
                merged.append(cur.copy())
        return merged

    @staticmethod
    def _mask_text(text: str, entities: List[Dict]) -> str:
        """對文字套用遮罩（從後往前替換以保持位置正確）"""
        if not entities:
            return text
        result = text
        for e in sorted(entities, key=lambda x: x["start_pos"], reverse=True):
            mask = MASK_MAP.get(e["entity_type"], f"[{e['entity_type']}]")
            s, end = e["start_pos"], e["end_pos"]
            if 0 <= s < end <= len(result):
                result = result[:s] + mask + result[end:]
        return result

    def _apply_masks_to_document(
        self,
        doc: Document,
        entities: List[Dict],
        full_text: str,
        output_dir: Path,
        job_id: str,
    ) -> Path:
        """套用遮罩並儲存 .docx / .json / .txt"""
        output_dir.mkdir(parents=True, exist_ok=True)

        masked_full = self._mask_text(full_text, entities)

        full_offset = 0
        for para in doc.paragraphs:
            para_text = para.text
            para_len = len(para_text)

            para_entities = [
                {
                    **e,
                    "start_pos": e["start_pos"] - full_offset,
                    "end_pos": e["end_pos"] - full_offset,
                }
                for e in entities
                if full_offset <= e["start_pos"] < full_offset + para_len
                and e["end_pos"] <= full_offset + para_len
            ]

            full_offset += para_len + 1

            if not para_entities or not para_text.strip():
                continue

            new_text = self._mask_text(para_text, para_entities)
            if new_text == para_text:
                continue

            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

        output_path = output_dir / f"{job_id}_deidentified.docx"
        doc.save(output_path)

        (output_dir / f"{job_id}_entities.json").write_text(
            json.dumps(entities, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        (output_dir / f"{job_id}_deidentified.txt").write_text(
            masked_full, encoding="utf-8"
        )

        logger.info(f"去識別化文件已儲存：{output_path}")
        return output_path


# ── 常見兩字中文詞（避免百家姓誤判） ────────────────────────────
_COMMON_TWO_CHAR_WORDS = {
    "方式", "方法", "方面", "方案", "方向",
    "王牌", "王者",
    "林業", "林地",
    "張貼",
    "高度", "高額", "高於",
    "黃金",
    "周年", "周期", "周知", "周圍",
    "何時", "何處", "何者",
    "許可", "許諾",
    "曾經",
    "簡稱", "簡述", "簡介", "簡易",
    "施行", "施工",
    "余額",
    "範圍", "範例", "範本",
    "丁方",
    "石材",
    "連續", "連帶", "連同",
    "田地",
    "馬上",
    "藍圖",
    "溫度",
    "龍頭",
}
