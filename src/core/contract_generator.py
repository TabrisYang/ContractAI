"""
合約生成引擎（方案 C：骨架拼裝 + LLM 局部填充）

流程：
1. LLM 萃取結構化需求（token 消耗低，只解析意圖）
2. 向量搜尋找最相似的真實合約骨架（不用 LLM）
3. 向量搜尋為每個需要的條款找最佳參考例句（不用 LLM）
4. LLM 只做「填充變數 + 客製化語句」，不從頭生成
5. 組合成完整合約，儲存 .docx
"""
import json
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import chromadb
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from loguru import logger
from sentence_transformers import SentenceTransformer

from .config import settings
from .llm.base import BaseLLMProvider, LLMMessage

CORPUS_COLLECTION = "contract_corpus"

# ── Prompt 設計（精簡，降低 token）──────────────────────────────────

EXTRACT_PROMPT = """從使用者的描述中萃取合約需求。只回傳 JSON，不含其他文字。

格式：
{
  "contract_type": "合約類型（如：服務合約、買賣合約、租賃合約、保密協議、委任合約、勞動合約）",
  "party_a": "甲方名稱（若未提及填 null）",
  "party_b": "乙方名稱（若未提及填 null）",
  "purpose": "合約主要目的（一句話）",
  "duration": "合約期間（若未提及填 null）",
  "amount": "金額（若未提及填 null）",
  "special_clauses": ["特殊條款需求1", "特殊條款需求2"],
  "missing_info": ["還需要確認的資訊1", "還需要確認的資訊2"]
}

若 missing_info 不為空，代表資訊不足，需要追問使用者。
若所有關鍵欄位（contract_type, party_a, party_b, purpose）都有值，missing_info 應為空陣列。"""

FILL_CLAUSE_PROMPT = """你是台灣法律合約起草專家。
根據以下資訊，將參考條款改寫為符合需求的條款文字。
只輸出條款內容本身，不加條款編號，使用繁體中文。

合約需求：
{requirements}

參考條款（從真實合約取出，請保留法律用語風格）：
{reference}

請將參考條款中的當事人名稱、金額、期間等替換為實際需求內容，保持法律措辭不變。"""

CLOSING_PROMPT = """根據以下合約需求，生成合約的「前言」和「簽署欄」，使用繁體中文。
只輸出這兩個部分，格式如下：

[前言]
（立合約書人...兩造當事人...）

[簽署欄]
（甲方簽章、乙方簽章的格式）

合約需求：{requirements}"""


# ── 資料結構 ──────────────────────────────────────────────────────

@dataclass
class GenerationRequirements:
    contract_type: str = ""
    party_a: str = ""
    party_b: str = ""
    purpose: str = ""
    duration: Optional[str] = None
    amount: Optional[str] = None
    special_clauses: list = field(default_factory=list)
    missing_info: list = field(default_factory=list)

    def is_complete(self) -> bool:
        return bool(
            self.contract_type and self.party_a and self.party_b and self.purpose
            and not self.missing_info
        )

    def to_summary(self) -> str:
        lines = [
            f"合約類型：{self.contract_type}",
            f"甲方：{self.party_a}",
            f"乙方：{self.party_b}",
            f"目的：{self.purpose}",
        ]
        if self.duration:
            lines.append(f"期間：{self.duration}")
        if self.amount:
            lines.append(f"金額：{self.amount}")
        if self.special_clauses:
            lines.append(f"特殊需求：{'、'.join(self.special_clauses)}")
        return "\n".join(lines)


@dataclass
class GenerationResult:
    gen_id: str
    stage: str          # "clarifying" | "draft_ready"
    question: str = ""  # 若 clarifying，追問的問題
    contract_text: str = ""  # 若 draft_ready，markdown 合約文字
    requirements: Optional[GenerationRequirements] = None


# ── 核心引擎 ──────────────────────────────────────────────────────

class ContractGenerator:

    # 標準合約條款類型（用於向量搜尋）
    STANDARD_CLAUSE_TYPES = [
        "服務範圍與工作內容",
        "合約期間與終止條件",
        "費用與付款方式",
        "保密義務",
        "智慧財產權歸屬",
        "違約責任與損害賠償",
        "爭議解決與管轄法院",
        "不可抗力條款",
    ]

    def __init__(self):
        self._embedder: Optional[SentenceTransformer] = None
        self._chroma: Optional[chromadb.ClientAPI] = None

    def _get_embedder(self) -> SentenceTransformer:
        if self._embedder is None:
            self._embedder = SentenceTransformer(settings.EMBEDDING_MODEL)
        return self._embedder

    def _get_collection(self):
        if self._chroma is None:
            self._chroma = chromadb.PersistentClient(path=str(settings.CHROMA_DIR))
        try:
            return self._chroma.get_collection(CORPUS_COLLECTION)
        except Exception:
            raise RuntimeError(
                "參考合約庫尚未建立索引。請先執行：python index_contracts.py"
            )

    def corpus_is_ready(self) -> bool:
        try:
            self._get_collection()
            return True
        except Exception:
            return False

    def corpus_info(self) -> dict:
        summary_path = settings.CHROMA_DIR / "corpus_summary.json"
        if summary_path.exists():
            return json.loads(summary_path.read_text())
        return {"total_contracts": 0, "total_chunks": 0}

    # ── 步驟 1：需求萃取（LLM，小 token）──────────────────────────

    async def extract_requirements(
        self,
        user_message: str,
        llm: BaseLLMProvider,
        history: list,
    ) -> GenerationRequirements:
        """從對話歷史萃取結構化需求，token 消耗最小化"""

        # 只傳最近 6 輪對話 + 萃取指令
        recent_history = history[-6:] if len(history) > 6 else history
        context = "\n".join(
            f"{'使用者' if m.get('role') == 'user' else '助手'}：{m.get('content', '')}"
            for m in recent_history
        )
        if context:
            prompt = f"對話紀錄：\n{context}\n\n最新訊息：{user_message}"
        else:
            prompt = user_message

        messages = [
            LLMMessage(role="system", content=EXTRACT_PROMPT),
            LLMMessage(role="user", content=prompt),
        ]
        raw = await llm.chat(messages, temperature=0.1, max_tokens=600)

        return self._parse_requirements(raw)

    def _parse_requirements(self, raw: str) -> GenerationRequirements:
        json_match = re.search(r'\{[\s\S]*\}', raw)
        if not json_match:
            return GenerationRequirements(
                missing_info=["請描述您需要的合約類型、甲乙方名稱及合約目的"]
            )
        try:
            data = json.loads(json_match.group())
        except json.JSONDecodeError:
            return GenerationRequirements(
                missing_info=["請描述您需要的合約類型、甲乙方名稱及合約目的"]
            )

        req = GenerationRequirements(
            contract_type=data.get("contract_type") or "",
            party_a=data.get("party_a") or "",
            party_b=data.get("party_b") or "",
            purpose=data.get("purpose") or "",
            duration=data.get("duration"),
            amount=data.get("amount"),
            special_clauses=data.get("special_clauses") or [],
            missing_info=data.get("missing_info") or [],
        )
        return req

    # ── 步驟 2：向量搜尋骨架 + 採用回饋重排（不用 LLM）──────────────

    def _search_reference_chunks(self, query: str, top_k: int = 3) -> list[dict]:
        """向量搜尋並以「cosine + β·採用分數」重排，回傳 [{id, text}]。

        採用分數用貝氏平滑：acceptance = (times_adopted + 1) / (times_retrieved + 2)，
        新範本樣本少時收斂到 0.5、累積夠多才逼近真實採用率。
        GEN_THOMPSON_SAMPLING 開啟時改從 Beta(adopted+1, refined+1) 抽樣，平衡探索/利用。
        並累計被選中段落的 times_retrieved。
        """
        collection = self._get_collection()
        embedder = self._get_embedder()
        q_emb = embedder.encode([query]).tolist()

        n_cand = max(top_k, settings.RAG_RERANK_CANDIDATES)
        results = collection.query(
            query_embeddings=q_emb,
            n_results=n_cand,
            include=["documents", "metadatas", "distances"],
        )
        docs = results.get("documents", [[]])[0]
        ids = results.get("ids", [[]])[0]
        metas = results.get("metadatas", [[]])[0]
        dists = results.get("distances", [[]])[0]

        beta = settings.GEN_ACCEPTANCE_WEIGHT
        use_thompson = settings.GEN_THOMPSON_SAMPLING
        ranked = []
        for i in range(len(docs)):
            cosine = 1.0 - float(dists[i]) if i < len(dists) else 0.0
            meta = metas[i] if i < len(metas) else {}
            retrieved = (meta or {}).get("times_retrieved", 0) or 0
            adopted = (meta or {}).get("times_adopted", 0) or 0
            refined = (meta or {}).get("times_refined_away", 0) or 0
            if use_thompson:
                acceptance = self._beta_sample(adopted + 1, refined + 1, seed=(i, retrieved))
            else:
                acceptance = (adopted + 1) / (retrieved + 2)
            cid = ids[i] if i < len(ids) else None
            ranked.append((cosine + beta * acceptance, docs[i], cid))

        ranked.sort(key=lambda x: x[0], reverse=True)
        top = ranked[:top_k]

        top_ids = [t[2] for t in top if t[2] is not None]
        self._bump_corpus_counts(collection, top_ids, field="times_retrieved")
        return [{"id": t[2], "text": t[1]} for t in top]

    def _search_reference_clauses(self, query: str, top_k: int = 3) -> list[str]:
        """相容舊介面：只回傳段落文字。"""
        return [c["text"] for c in self._search_reference_chunks(query, top_k)]

    @staticmethod
    def _beta_sample(alpha: float, beta_param: float, seed=None) -> float:
        """Beta 分布抽樣（Thompson Sampling）。以 numpy 預設亂數產生器。"""
        import numpy as np
        rng = np.random.default_rng()
        return float(rng.beta(alpha, beta_param))

    @staticmethod
    def _bump_corpus_counts(collection, chunk_ids: list, field: str, amount: int = 1):
        """為 corpus 段落 metadata 計數器 +amount（times_retrieved/adopted/refined_away）。"""
        if not chunk_ids:
            return
        try:
            cur = collection.get(ids=chunk_ids, include=["metadatas"])
            got_ids = cur.get("ids") or []
            metas = cur.get("metadatas") or []
            new_metas = []
            for m in metas:
                m = dict(m or {})
                m[field] = (m.get(field, 0) or 0) + amount
                new_metas.append(m)
            if got_ids:
                collection.update(ids=got_ids, metadatas=new_metas)
        except Exception as e:
            logger.warning(f"更新 corpus 段落計數失敗（{field}）：{e}")

    # ── 步驟 3：逐條填充（LLM，每次只處理一條）──────────────────

    async def _fill_clause(
        self,
        clause_type: str,
        requirements: GenerationRequirements,
        llm: BaseLLMProvider,
    ) -> tuple[str, list]:
        """為單一條款找參考例句並讓 LLM 客製化填充，回傳 (條款文字, 參考段落 id)。"""
        # 搜尋此條款類型的參考例句（純向量搜尋 + 採用回饋重排，不用 LLM）
        query = f"{requirements.contract_type} {clause_type}"
        ref_chunks = self._search_reference_chunks(query, top_k=2)
        references = [c["text"] for c in ref_chunks]
        used_ids = [c["id"] for c in ref_chunks if c["id"]]
        ref_text = "\n---\n".join(references) if references else "（無參考例句）"

        prompt = FILL_CLAUSE_PROMPT.format(
            requirements=requirements.to_summary(),
            reference=ref_text,
        )
        messages = [LLMMessage(role="user", content=prompt)]
        result = await llm.chat(messages, temperature=0.2, max_tokens=400)
        return result.strip(), used_ids

    # ── 步驟 4：生成前言與簽署欄（LLM，小 token）────────────────

    async def _generate_preamble_and_closing(
        self,
        requirements: GenerationRequirements,
        llm: BaseLLMProvider,
    ) -> tuple[str, str]:
        prompt = CLOSING_PROMPT.format(requirements=requirements.to_summary())
        messages = [LLMMessage(role="user", content=prompt)]
        raw = await llm.chat(messages, temperature=0.2, max_tokens=600)

        # 解析前言與簽署欄
        preamble = ""
        closing = ""
        preamble_match = re.search(r'\[前言\]\s*([\s\S]*?)(?=\[簽署欄\]|$)', raw)
        closing_match = re.search(r'\[簽署欄\]\s*([\s\S]*?)$', raw)
        if preamble_match:
            preamble = preamble_match.group(1).strip()
        if closing_match:
            closing = closing_match.group(1).strip()
        return preamble, closing

    # ── 主流程 ──────────────────────────────────────────────────────

    async def generate(
        self,
        requirements: GenerationRequirements,
        llm: BaseLLMProvider,
        gen_id: str,
        output_dir: Path,
    ) -> str:
        """
        方案 C 核心：骨架拼裝 + LLM 局部填充
        回傳 markdown 格式的合約文字
        """
        output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"開始生成合約 {gen_id}，類型：{requirements.contract_type}")

        # 決定本合約需要哪些條款
        clause_types = self._select_clause_types(requirements)

        # 並行生成各條款（但逐條送 LLM，每次 token 量小）
        sections = []

        # 前言
        preamble, closing = await self._generate_preamble_and_closing(requirements, llm)

        # 逐條生成
        used_chunk_ids: list = []
        for i, clause_type in enumerate(clause_types, 1):
            logger.info(f"  生成第 {i} 條：{clause_type}")
            content, used_ids = await self._fill_clause(clause_type, requirements, llm)
            used_chunk_ids.extend(used_ids)
            sections.append((i, clause_type, content))

        # 記錄本次生成參考了哪些 corpus 段落（供採用/refine 回饋更新權重）
        unique_ids = sorted(set(used_chunk_ids))
        (output_dir / f"{gen_id}_used_chunks.json").write_text(
            json.dumps(unique_ids, ensure_ascii=False), encoding="utf-8"
        )

        # 組合成完整合約（markdown）
        contract_md = self._assemble_markdown(requirements, preamble, sections, closing)

        # 儲存 markdown
        md_path = output_dir / f"{gen_id}_draft.md"
        md_path.write_text(contract_md, encoding="utf-8")

        # 生成 .docx
        docx_path = output_dir / f"{gen_id}_draft.docx"
        self._save_docx(requirements, preamble, sections, closing, docx_path)

        logger.info(f"合約生成完成：{docx_path}")
        return contract_md

    async def refine(
        self,
        gen_id: str,
        feedback: str,
        llm: BaseLLMProvider,
        output_dir: Path,
    ) -> str:
        """根據使用者回饋修改現有草稿"""
        md_path = output_dir / f"{gen_id}_draft.md"
        if not md_path.exists():
            raise FileNotFoundError(f"找不到草稿：{md_path}")

        current_text = md_path.read_text(encoding="utf-8")

        messages = [
            LLMMessage(
                role="system",
                content="你是台灣法律合約起草專家。根據使用者的修改指示，修改合約內容。只輸出修改後的完整合約，使用繁體中文，不加任何說明。",
            ),
            LLMMessage(
                role="user",
                content=f"目前合約：\n\n{current_text[:4000]}\n\n修改指示：{feedback}",
            ),
        ]
        refined = await llm.chat(messages, temperature=0.2, max_tokens=4000)

        # 更新儲存
        md_path.write_text(refined, encoding="utf-8")

        # 重新生成 docx（簡化版：直接存純文字 docx）
        docx_path = output_dir / f"{gen_id}_draft.docx"
        self._save_plain_docx(refined, docx_path)

        # 回饋：本次草稿被要求修改 → 參考段落記為「被 refine」（負訊號）
        self.record_generation_feedback(gen_id, adopted=False, output_dir=output_dir, feedback=feedback)

        return refined

    def _used_chunk_ids(self, gen_id: str, output_dir: Path) -> list:
        path = output_dir / f"{gen_id}_used_chunks.json"
        if not path.exists():
            return []
        try:
            return json.loads(path.read_text(encoding="utf-8")) or []
        except Exception:
            return []

    def record_generation_feedback(
        self,
        gen_id: str,
        adopted: bool,
        output_dir: Path,
        feedback: str = "",
    ) -> int:
        """記錄生成回饋並更新參考段落權重（回饋迴路 C）。

        adopted=True（使用者下載/保存最終合約）→ times_adopted +1（提權）
        adopted=False（被 refine）→ times_refined_away +1（降權）
        回傳更新的段落數。同時寫入統一回饋庫（loop="gen"）。
        """
        from .feedback_store import feedback_store
        chunk_ids = self._used_chunk_ids(gen_id, output_dir)
        feedback_store.record(
            loop="gen",
            job_id=gen_id,
            target_ref={"chunk_ids": chunk_ids},
            signal={"adopted": adopted, "feedback": feedback},
        )
        if not chunk_ids:
            return 0
        field = "times_adopted" if adopted else "times_refined_away"
        try:
            collection = self._get_collection()
        except Exception:
            return 0
        self._bump_corpus_counts(collection, chunk_ids, field=field)
        return len(chunk_ids)

    # ── 工具方法 ──────────────────────────────────────────────────

    def _select_clause_types(self, req: GenerationRequirements) -> list[str]:
        """根據合約類型選擇適合的條款"""
        base = [
            "服務範圍與工作內容",
            "合約期間與終止條件",
            "費用與付款方式",
            "違約責任與損害賠償",
            "爭議解決與管轄法院",
        ]
        # 根據類型加入特殊條款
        ct = req.contract_type
        if "保密" in ct or any("保密" in s for s in req.special_clauses):
            base.insert(3, "保密義務")
        if "軟體" in ct or "開發" in ct or "設計" in ct:
            base.insert(4, "智慧財產權歸屬")
        if "租賃" in ct:
            base = ["租賃物說明", "租金與付款方式", "使用限制", "修繕責任",
                    "合約期間與終止條件", "違約責任與損害賠償", "爭議解決與管轄法院"]
        if "勞動" in ct or "僱傭" in ct:
            base = ["職務內容", "薪資與福利", "工作時間", "保密義務",
                    "競業禁止", "合約期間與終止條件", "違約責任與損害賠償"]
        # 加入使用者特殊需求
        for special in req.special_clauses:
            if special not in base:
                base.append(special)
        return base

    def _assemble_markdown(self, req, preamble, sections, closing) -> str:
        lines = [
            f"# {req.contract_type}",
            "",
            "---",
            "",
            preamble,
            "",
        ]
        for i, clause_type, content in sections:
            lines.append(f"## 第{_num_to_chinese(i)}條　{clause_type}")
            lines.append("")
            lines.append(content)
            lines.append("")
        lines += ["---", "", closing, ""]
        return "\n".join(lines)

    def _save_docx(self, req, preamble, sections, closing, path: Path):
        """生成格式化的 .docx"""
        doc = Document()

        # 標題
        title = doc.add_heading(req.contract_type, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # 前言
        if preamble:
            para = doc.add_paragraph(preamble)
            para.paragraph_format.first_line_indent = Pt(24)

        doc.add_paragraph()

        # 各條款
        for i, clause_type, content in sections:
            heading = doc.add_heading(f"第{_num_to_chinese(i)}條　{clause_type}", level=2)
            para = doc.add_paragraph(content)
            para.paragraph_format.first_line_indent = Pt(24)
            doc.add_paragraph()

        # 簽署欄
        doc.add_paragraph("─" * 40)
        if closing:
            doc.add_paragraph(closing)

        doc.save(path)

    def _save_plain_docx(self, text: str, path: Path):
        """將 markdown 文字儲存為簡單 .docx"""
        doc = Document()
        for line in text.splitlines():
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=0)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("---"):
                doc.add_paragraph("─" * 40)
            else:
                doc.add_paragraph(line)
        doc.save(path)


def _num_to_chinese(n: int) -> str:
    mapping = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
               6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
               11: "十一", 12: "十二", 13: "十三", 14: "十四", 15: "十五"}
    return mapping.get(n, str(n))


# ── 全域單例 ──────────────────────────────────────────────────────

_generator: Optional[ContractGenerator] = None


def get_generator() -> ContractGenerator:
    global _generator
    if _generator is None:
        _generator = ContractGenerator()
    return _generator
